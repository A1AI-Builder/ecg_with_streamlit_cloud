import os
from datetime import datetime
import streamlit as st
from io import BytesIO
from PIL import Image
from docx import Document
from docx.shared import Inches
import google.generativeai as genai
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()

# Configure the generative model
api_key = os.getenv("GOOGLE_API_KEY")
if api_key is None:
    raise ValueError("GEMINI_API_KEY is not set in environment variables")
genai.configure(api_key=api_key)

# Set up the model
generation_config = {
    "temperature": 0.8,
    "top_p": 0.9,
    "top_k": 50,
    "max_output_tokens": 2048,
    "response_mime_type": "text/plain",
}

model = genai.GenerativeModel(
    model_name="gemini-1.5-pro",
    generation_config=generation_config,
)

@st.cache_data
def cached_generate_ecg_details(image_bytes):
    """Cached function to generate ECG details from the image."""
    image = Image.open(BytesIO(image_bytes))
    prompt = f"""Analyze this ECG image and provide a detailed report. Fill in ALL fields based on the information you can extract from the image. If you absolutely cannot determine a piece of information, state 'Unable to determine from the provided ECG image.' ..."""  # Shortened for brevity.
    chat_session = model.start_chat(history=[])
    response = chat_session.send_message([prompt, image])
    return response.text

def create_doc(report_text, ecg_image_bytes):
    """Create a Word document with the ECG report and image."""
    doc = Document()
    doc.add_heading('ECG ANALYSIS REPORT', 0)
    for line in report_text.split("\n"):
        if line.strip() == '':
            continue
        if line.startswith('**') and line.endswith('**'):
            doc.add_heading(line.strip('**'), level=1)
        elif line.startswith('-'):
            doc.add_paragraph(line.strip(), style='List Bullet')
        else:
            doc.add_paragraph(line.strip())

    doc.add_heading('ECG Tracing:', level=1)
    image_stream = BytesIO(ecg_image_bytes)
    doc.add_picture(image_stream, width=Inches(6))

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def main():
    st.title("Electrocardiogram Analysis Insights into Heart Health")
    st.header("Attachments")
    ecg_image = st.file_uploader("Upload ECG Image", type=["png", "jpg", "jpeg"])

    if ecg_image is not None:
        # Resize and display the uploaded image
        image = Image.open(ecg_image)
        resized_image = image.resize((800, 600), Image.Resampling.LANCZOS)
        image_bytes = BytesIO()
        resized_image.save(image_bytes, format="PNG")
        st.image(resized_image, caption='Uploaded ECG Image', use_column_width=True)

        if st.button("Generate ECG Report"):
            with st.spinner("Analyzing ECG image..."):
                ecg_details = cached_generate_ecg_details(image_bytes.getvalue())
            st.header("Generated ECG Report")
            st.markdown(ecg_details)

            # Store the generated report in session state
            st.session_state.ecg_details = ecg_details

        # Check if report has been generated
        if hasattr(st.session_state, 'ecg_details'):
            doc_file_stream = create_doc(st.session_state.ecg_details, image_bytes.getvalue())
            st.download_button(
                label="Download ECG Report",
                data=doc_file_stream,
                file_name="ECG_Report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

if __name__ == '__main__':
    main()
